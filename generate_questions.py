from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Dark cyber colors
DARK_BG = RGBColor(0x0A, 0x0A, 0x0F)      # #0A0A0F
ACCENT = RGBColor(0x3A, 0x86, 0xFF)        # #3A86FF
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
DARK_GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN_ACCENT = RGBColor(0x00, 0xCC, 0x88)
RED_ACCENT = RGBColor(0xFF, 0x44, 0x44)
ORANGE = RGBColor(0xFF, 0x99, 0x00)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_colored_paragraph(cell, text, color=WHITE, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Calibri'
    return p


def add_accent_heading(text, level=1):
    """Add a colored heading (blue accent)."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level <= 2 else RGBColor(0xDD, 0xDD, 0xDD)
    return h


def add_qa_block(q_num, question, answer, section_name=""):
    """Add a question-answer block with formatting."""
    # Question row
    q_table = doc.add_table(rows=1, cols=1)
    q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    q_cell = q_table.cell(0, 0)
    set_cell_shading(q_cell, "1A1A2E")
    q_text = f"Вопрос {q_num}: {question}"
    add_colored_paragraph(q_cell, q_text, color=ACCENT, bold=True, size=12)

    # Answer row
    a_table = doc.add_table(rows=1, cols=1)
    a_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    a_cell = a_table.cell(0, 0)
    set_cell_shading(a_cell, "0D0D1A")
    add_colored_paragraph(a_cell, answer, color=LIGHT_GRAY, size=11)

    # Spacer
    doc.add_paragraph("")

    # Set borders for both tables
    for table in [q_table, a_table]:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)


# ══════════════════════════════════════════════════════════════
#                     TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ВОПРОСЫ И ОТВЕТЫ")
run.font.size = Pt(36)
run.font.color.rgb = ACCENT
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("К ЗАЩИТЕ ДИПЛОМНОГО ПРОЕКТА")
run.font.size = Pt(20)
run.font.color.rgb = WHITE

doc.add_paragraph("")

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    "Аудиоплеер (Desktop + Android) • Интернет-магазин мерча (Web + Mobile)\n"
    "Технологии: C# / Kotlin / Python / TypeScript / React / React Native"
)
run.font.size = Pt(13)
run.font.color.rgb = LIGHT_GRAY

doc.add_paragraph("")
doc.add_paragraph("")

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f"2025 — {datetime.datetime.now().year}")
run.font.size = Pt(14)
run.font.color.rgb = DARK_GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                   TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
add_accent_heading("Содержание", level=1)

toc_items = [
    ("Раздел 1", "Общие вопросы о проекте", "1–7"),
    ("Раздел 2", "Аудиоплеер: Десктопная версия (C# / Avalonia)", "8–19"),
    ("Раздел 3", "Аудиоплеер: Android-версия (Kotlin / Jetpack Compose)", "20–29"),
    ("Раздел 4", "Интернет-магазин: Бэкенд (Django REST)", "30–40"),
    ("Раздел 5", "Интернет-магазин: Веб-фронтенд (React / TypeScript)", "41–49"),
    ("Раздел 6", "Интернет-магазин: Мобильное приложение (React Native)", "50–54"),
    ("Раздел 7", "База данных и инфраструктура", "55–61"),
    ("Раздел 8", "Архитектура, безопасность и тестирование", "62–69"),
    ("Раздел 9", "Сравнение платформ и перспективы развития", "70–75"),
]

toc_table = doc.add_table(rows=1, cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header row
for i, hdr_text in enumerate(["Раздел", "Название", "Вопросы"]):
    cell = toc_table.cell(0, i)
    set_cell_shading(cell, "1A1A2E")
    add_colored_paragraph(cell, hdr_text, color=ACCENT, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

for sec, name, nums in toc_items:
    row = toc_table.add_row()
    for i, val in enumerate([sec, name, nums]):
        cell = row.cells[i]
        set_cell_shading(cell, "0A0A0F")
        c = ACCENT if i == 0 else (LIGHT_GRAY if i == 1 else DARK_GRAY)
        add_colored_paragraph(cell, val, color=c, size=10, align=WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT)

# Set borders for TOC table
tbl = toc_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="3A86FF"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#                   SECTION DIVIDER
# ══════════════════════════════════════════════════════════════
def add_section_header(title_text, number):
    """Add a decorative section header."""
    h = doc.add_heading(f"{number}. {title_text}", level=1)
    for run in h.runs:
        run.font.color.rgb = ACCENT
        run.font.size = Pt(18)
    # Underline
    line = doc.add_paragraph()
    line_run = line.add_run("─" * 80)
    line_run.font.color.rgb = ACCENT
    line_run.font.size = Pt(6)


# ══════════════════════════════════════════════════════════════
#                    QUESTIONS & ANSWERS
# ══════════════════════════════════════════════════════════════

# ──────────────────────────────────────────────
# SECTION 1: Общие вопросы
# ──────────────────────────────────────────────
add_section_header("Общие вопросы о проекте", "Раздел 1")

add_qa_block(1,
    "Какова цель и задачи вашего дипломного проекта?",
    "Цель проекта — разработка跨платформенного аудиоплеера с расширенным функционалом и торговой "
    "интернет-площадки по продаже мерча российских рэп-исполнителей. Задачи проекта:\n"
    "1) Разработка десктопного аудиоплеера на C# с использованием Avalonia UI для Windows;\n"
    "2) Разработка мобильного аудиоплеера на Kotlin с Jetpack Compose для Android;\n"
    "3) Создание REST API на Django REST Framework для интернет-магазина стикеров и мерча;\n"
    "4) Разработка веб-клиента на React + TypeScript;\n"
    "5) Разработка мобильного клиента на React Native;\n"
    "6) Интеграция JWT-аутентификации, корзины, оформления заказов;\n"
    "7) Создание кастомной админ-панели для управления контентом;\n"
    "8) Миграция базы данных с PostgreSQL на SQLite для упрощения развёртывания."
)

add_qa_block(2,
    "Какие технологии и почему вы выбрали для каждой части проекта?",
    "Выбор технологий обоснован спецификой каждой платформы:\n"
    "• Десктопный плеер: C# / .NET 10 + Avalonia UI 11. Выбор обусловлен кроссплатформенностью Avalonia "
    "(Windows, macOS, Linux), современным XAML-движком с аппаратным ускорением и встроенной поддержкой "
    "Fluent Design с акриловым размытием. NAudio 2.2.1 обеспечивает низкоуровневый доступ к аудиопотоку "
    "(WaveOutEvent, AudioFileReader) для программного эквалайзера.\n"
    "• Android-плеер: Kotlin + Jetpack Compose + AndroidX Media3. Compose — современный декларативный "
    "UI-фреймворк от Google, пришедший на смену XML-разметке. Media3 (ExoPlayer) — стандарт для "
    "аудио/видео на Android с поддержкой MediaSessionService для фонового воспроизведения.\n"
    "• Бэкенд: Django 5.2 + DRF. Django предоставляет готовую админ-панель, ORM, миграции и "
    "аутентификацию «из коробки». DRF — стандарт для создания REST API на Python.\n"
    "• Веб-фронтенд: React 18 + Vite + TypeScript. React — самая популярная библиотека для SPA, "
    "Vite обеспечивает быструю сборку и HMR, TypeScript добавляет статическую типизацию.\n"
    "• Мобильное приложение: React Native (Expo SDK 56). React Native позволяет использовать общие "
    "компоненты и логику с веб-версией, сокращая время разработки.\n"
    "• База данных: PostgreSQL (разработка) / SQLite (финальная версия). PostgreSQL — надёжная "
    "промышленная БД, SQLite — файловая БД без необходимости отдельного сервера."
)

add_qa_block(3,
    "Какова общая архитектура проекта?",
    "Проект состоит из двух независимых направлений, объединённых общей тематикой (музыка) и "
    "стилистикой (киберпанк, тёмная тема, неоновая цветовая палитра #0A0A0F / #3A86FF):\n\n"
    "Направление 1 — STREET PLAYER (Аудиоплеер):\n"
    "• Десктопная версия: C# / .NET 10, Avalonia UI, MVVM через CommunityToolkit.Mvvm, "
    "NAudio для воспроизведения, TagLibSharp для ID3-тегов и обложек.\n"
    "• Android-версия: Kotlin, Jetpack Compose, AndroidX Media3 (ExoPlayer), MediaSessionService "
    "для фонового воспроизведения, MediaStore для сканирования музыки.\n"
    "• Тесты: xUnit (18 unit-тестов для ViewModel).\n\n"
    "Направление 2 — STICKER MERCH (Интернет-магазин):\n"
    "• Бэкенд: Django 5.2 + DRF (9 REST-эндпоинтов), JWT-аутентификация (simplejwt), "
    "две админ-панели (стандартная /admin/ и кастомная /cp/ с 16 view).\n"
    "• Веб-фронтенд: React 18 + Vite + TypeScript (9 страниц, 1688 строк CSS, 53 SVG-иконки, "
    "10 иконок артистов).\n"
    "• Мобильное приложение: React Native (Expo SDK 56) с кастомной стековой навигацией, "
    "анимированным боковым меню и мок-данными в качестве fallback.\n"
    "• База данных: PostgreSQL → SQLite (миграция выполнена, все 243 объекта сохранены)."
)

add_qa_block(4,
    "Почему вы решили объединить плеер и магазин в одном проекте?",
    "Формально это две независимые подсистемы, объединённые на уровне дипломного проекта "
    "общей музыкальной тематикой. Плеер ориентирован на прослушивание музыки, а магазин — "
    "на приобретение мерча любимых исполнителей. С архитектурной точки зрения они не связаны "
    "(нет общей логики, API или базы данных), но вместе демонстрируют широкий спектр навыков "
    "разработки: десктопные приложения (C#), мобильные native-приложения (Kotlin), "
    "веб-бэкенд (Python), веб-фронтенд (TypeScript) и кроссплатформенные мобильные "
    "приложения (React Native). Это позволяет показать владение пятью языками программирования "
    "и четырьмя различными платформами."
)

add_qa_block(5,
    "Какие паттерны проектирования вы использовали?",
    "В проекте применяются следующие паттерны:\n"
    "• MVVM (Model-View-ViewModel) — в десктопном плеере через CommunityToolkit.Mvvm. "
    "View (MainWindow.axaml) не содержит логики — она вынесена в PlayerViewModel с "
    "атрибутами [ObservableProperty] и [RelayCommand] для source-генерации кода.\n"
    "• MVVM — в Android-плеере через AndroidViewModel + StateFlow. UI (PlayerScreen.kt) "
    "наблюдает за состоянием через StateFlow, ViewModel управляет бизнес-логикой.\n"
    "• MVC (Model-View-Controller) — на бэкенде Django: Model (ORM), View (APIView), "
    "Template/Serializer (представление).\n"
    "• Singleton — Django ORM и пул соединений с БД.\n"
    "• Service Layer — PlaybackService на Android как MediaSessionService.\n"
    "• Repository — MusicScanner инкапсулирует логику доступа к MediaStore.\n"
    "• Data Mapper — Django ORM с managed=False для legacy-таблиц.\n"
    "• Observer — StateFlow в Android ViewModel, INotifyPropertyChanged в десктопе."
)

add_qa_block(6,
    "С какими основными трудностями вы столкнулись?",
    "Основные трудности:\n"
    "1) Кроссплатформенность десктопного плеера — Avalonia UI, будучи кроссплатформенным "
    "фреймворком, имеет отличия в отрисовке окон. Пришлось адаптировать разметку под Windows.\n"
    "2) Фоновое воспроизведение на Android — начиная с Android 13, Google ужесточила требования "
    "к foreground-сервисам. Обязателен вызов startForeground() с уведомлением, иначе — "
    "IllegalStateException.\n"
    "3) Проблемы с кириллицей в пути — из-за того, что проект находится в папке «E:\\илья проект», "
    "Android Gradle Plugin выдавал ошибку. Решение — android.overridePathCheck=true.\n"
    "4) Несовместимость JDK — AGP 8.2.2 работает только с JDK 17, JDK 24 выдавал ошибки.\n"
    "5) Миграция PostgreSQL → SQLite — необходимость переноса данных через dump/restore с "
    "правильной кодировкой UTF-8.\n"
    "6) Обработка managed=False моделей — Django не управляет миграциями для legacy-таблиц, "
    "пришлось временно включать managed=True для создания таблиц, затем возвращать обратно."
)

add_qa_block(7,
    "Какие языки программирования и сколько строк кода в проекте?",
    "Проект написан на 5 языках программирования (оценка приблизительная):\n"
    "• C#: ~1200 строк (десктопный плеер + тесты)\n"
    "• Kotlin: ~1100 строк (Android-плеер)\n"
    "• Python: ~1500 строк (бэкенд Django + скрипты)\n"
    "• TypeScript: ~3300 строк (фронтенд React + React Native)\n"
    "• CSS: ~1700 строк (дизайн-система в едином файле index.css)\n"
    "• XAML: ~200 строк (десктопный UI)\n"
    "• SQL: ~250 строк (seed-скрипт с сырыми запросами)\n"
    "Общее количество: около 9000-10000 строк кода."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 2: Desktop Player
# ──────────────────────────────────────────────
add_section_header("Аудиоплеер: Десктопная версия (C# / Avalonia)", "Раздел 2")

add_qa_block(8,
    "Почему вы выбрали Avalonia UI вместо WPF или WinForms?",
    "Avalonia UI выбран по нескольким причинам:\n"
    "1) Кроссплатформенность — Avalonia работает на Windows, macOS и Linux, в то время как WPF "
    "привязан только к Windows. Это позволяет в будущем запустить плеер на других ОС без "
    "переписывания кода.\n"
    "2) Современный XAML-движок с аппаратным ускорением через Skia — обеспечивает плавный "
    "рендеринг независимо от разрешения экрана.\n"
    "3) Встроенная поддержка Fluent Design — AcrylicBlur (акриловое размытие), тёмная тема, "
    "кастомные стили.\n"
    "4) Совместимость с .NET 10 — последняя версия платформы с AOT-компиляцией и улучшенной "
    "производительностью.\n"
    "5) MVVM-friendly — CommunityToolkit.Mvvm разработан специально для Avalonia и WPF, "
    "предоставляя source-генераторы для ObservableProperty и RelayCommand.\n\n"
    "По сравнению с WinForms: Avalonia более современный, декларативный, с аппаратным ускорением "
    "и поддержкой биндингов. По сравнению с WPF: Avalonia кроссплатформенный и не требует .NET Framework."
)

add_qa_block(9,
    "Расскажите подробнее об архитектуре MVVM в десктопном плеере.",
    "Архитектура MVVM реализована через библиотеку CommunityToolkit.Mvvm 8.3.2 с использованием "
    "source-генераторов (атрибуты вместо ручного кода):\n\n"
    "• Model: PlaylistItem.cs — простая модель трека (18 строк) с полями Path, Title, Artist, "
    "Album, Duration; реализует INotifyPropertyChanged для UI-биндинга.\n\n"
    "• ViewModel: PlayerViewModel.cs (488 строк) — наследуется от ObservableObject. "
    "Поля, помеченные [ObservableProperty], автоматически генерируют свойства с уведомлением "
    "об изменении. Методы с [RelayCommand] генерируют команды для привязки из XAML. "
    "Основные компоненты ViewModel:\n"
    "  — Playlist: ObservableCollection<PlaylistItem> — текущий плейлист;\n"
    "  — CurrentTrack: PlaylistItem? — текущий трек;\n"
    "  — IsPlaying, CurrentPosition, Duration, Volume — состояние плеера;\n"
    "  — SearchQuery — строка поиска с фильтрацией через LINQ;\n"
    "  — IsRepeat, IsAutoPlay, IsShuffle — режимы воспроизведения.\n\n"
    "• View: MainWindow.axaml (197 строк XAML) — декларативное описание интерфейса с "
    "привязками {Binding CurrentTrack.Title}, {Binding IsPlaying, Converter=...} и "
    "командами {Binding PlayPauseCommand}.\n\n"
    "Code-behind (MainWindow.axaml.cs, 30 строк) — минимален: создаёт ViewModel, "
    "устанавливает DataContext, подписывается на событие Closing для сохранения плейлиста."
)

add_qa_block(10,
    "Как реализован эквалайзер и какие частоты он регулирует?",
    "Эквалайзер реализован в классе EqualizerSampleProvider.cs (54 строки) через NAudio "
    "и BiQuadFilter (биквадратный фильтр, стандартный DSP-алгоритм в аудиообработке):\n\n"
    "• Фильтр низких частот (LowShelf) на 80 Гц с усилением/ослаблением +/-24 дБ — "
    "регулирует бас.\n"
    "• Фильтр высоких частот (HighShelf) на 6000 Гц с усилением/ослаблением +/-24 дБ — "
    "регулирует высокие частоты / «звонкость».\n\n"
    "EqualizerSampleProvider оборачивает AudioFileReader и применяет фильтры к аудиопотоку "
    "«на лету» через переопределённый метод Read(). Пользователь может независимо регулировать "
    "уровень басов и высоких частот через слайдеры в UI.\n\n"
    "Альтернативные подходы, которые рассматривались: использование готовых DSP-эффектов "
    "NAudio (WaveFilter) и написание графического 10-полосного эквалайзера. Выбран "
    "двухполосный вариант как оптимальный по соотношению сложность/функциональность."
)

add_qa_block(11,
    "Как работает сканирование музыкальных файлов на десктопе?",
    "Сканирование реализовано в PlayerViewModel через два метода:\n\n"
    "1) ScanDrives() — сканирует все доступные диски (фиксированные + съёмные) рекурсивно "
    "на глубину до 4 уровней вложенности. Для каждого найденного аудиофайла создаётся "
    "PlaylistItem с чтением ID3-тегов через TagLibSharp.\n\n"
    "2) SelectFolder() / AddFiles() — пользователь может вручную выбрать папку через "
    "диалог выбора папки или отдельные файлы.\n\n"
    "Дедупликация: перед добавлением каждого файла проверяется, нет ли его уже в плейлисте "
    "(по полному пути Path). Это предотвращает дублирование при повторном сканировании.\n\n"
    "Поддерживаемые форматы: mp3, wav, flac, ogg, aac, wma, m4a. "
    "Фильтрация: исключаются системные папки (Windows, Program Files и т.д.)."
)

add_qa_block(12,
    "Как реализовано сохранение и загрузка плейлиста?",
    "Плейлист сохраняется в формате JSON в папку %APPDATA%\\STREET PLAYER\\playlist.json. "
    "Это обеспечивает:\n"
    "• Сохранение между сессиями — плейлист автоматически загружается при старте;\n"
    "• Переносимость — файл можно скопировать на другой компьютер;\n"
    "• Человекочитаемость — можно просмотреть/отредактировать в любом текстовом редакторе.\n\n"
    "Сериализация: для каждого PlaylistItem сохраняются поля Path, Title, Artist, Album, Duration. "
    "Используется System.Text.Json (встроенный в .NET 10 сериализатор).\n\n"
    "Автосохранение: триггером является событие Closing главного окна — перед закрытием "
    "вызывается SavePlaylistCommand. Также есть ручное сохранение через SavePlaylistCommand "
    "в меню.\n\n"
    "При загрузке проверяется существование каждого файла — если файл удалён или перемещён, "
    "он пропускается с выводом предупреждения."
)

add_qa_block(13,
    "Какие форматы аудио поддерживаются и как читаются ID3-теги?",
    "Поддерживаются форматы: mp3, wav, flac, ogg, aac, wma, m4a.\n\n"
    "Чтение тегов реализовано через библиотеку TagLibSharp 2.3.0:\n"
    "• Извлечение Title, Artist, Album из ID3v2-тегов (mp3), Vorbis Comments (flac, ogg) "
    "и соответствующих метаданных для других форматов;\n"
    "• Чтение обложки альбома (APIC-фрейм для mp3, picture block для flac);\n"
    "• Длительность трека вычисляется через AudioFileReader.Duration (NAudio) или "
    "через TagLibSharp (как fallback).\n\n"
    "Если ID3-теги отсутствуют или повреждены, название файла используется как имя трека. "
    "Обложка в этом случае не отображается (показывается заглушка в виде иконки музыки)."
)

add_qa_block(14,
    "Как работает поиск по плейлисту?",
    "Поиск реализован через поле SearchQuery во ViewModel с немедленной фильтрацией "
    "(real-time):\n\n"
    "• При каждом изменении SearchQuery (биндинг в XAML через TextBox с UpdateSourceTrigger=PropertyChanged) "
    "вызывается метод ApplyFilter().\n"
    "• Фильтрация: PlaylistItem, у которых Title или Path содержат SearchQuery "
    "(регистронезависимое сравнение, StringComparison.OrdinalIgnoreCase).\n"
    "• Отфильтрованный список отображается в ListBox через CollectionViewSource "
    "с использованием BindingListCollectionView для поддержки поиска без полной "
    "перезагрузки списка.\n"
    "• UI: TextBox для ввода запроса в верхней части окна, результаты обновляются "
    "по мере ввода без задержки."
)

add_qa_block(15,
    "Расскажите о режимах воспроизведения: Repeat, AutoPlay, Shuffle.",
    "Режимы воспроизведения реализованы как отдельные boolean-свойства во ViewModel:\n\n"
    "• Repeat (IsRepeat): при окончании текущего трека он запускается сначала. "
    "Используется для зацикленного прослушивания одного трека.\n\n"
    "• AutoPlay (IsAutoPlay): после завершения текущего трека автоматически "
    "запускается следующий по порядку в плейлисте. Это стандартный режим «играть всё подряд».\n\n"
    "• Shuffle (IsShuffle): при включении перемешивает плейлист с использованием "
    "алгоритма Fisher-Yates (Knuth shuffle). После завершения трека выбирается "
    "случайный следующий. При выключении Shuffle возвращается исходный порядок.\n\n"
    "Приоритет режимов: Repeat > Shuffle > AutoPlay. Если все три выключены, "
    "воспроизведение останавливается после окончания текущего трека."
)

add_qa_block(16,
    "Как NAudio взаимодействует с ViewModel?",
    "NAudio используется через PlayerViewModel следующим образом:\n\n"
    "1) Инициализация: при выборе трека создаётся AudioFileReader (читает файл с диска).\n"
    "2) Обёртка: AudioFileReader оборачивается в EqualizerSampleProvider (эквалайзер).\n"
    "3) Воспроизведение: WaveOutEvent (аудиовыход Windows) инициализируется "
    "равелизатором и запускается через Play().\n"
    "4) Таймер позиции: Timer с интервалом ~250 мс опрашивает CurrentPosition и "
    "обновляет свойство CurrentPosition в ViewModel (что обновляет Slider в UI).\n"
    "5) Управление: Play/Pause через waveOut.Play() / waveOut.Pause();\n"
    "   Остановка: waveOut.Stop() с последующим Dispose();\n"
    "   Громкость: waveOut.Volume (0.0 – 1.0);\n"
    "   Перемотка: reader.CurrentTime = TimeSpan.FromSeconds(position).\n\n"
    "ViewModel инкапсулирует весь lifecycle NAudio: создание, воспроизведение, "
    "остановку и освобождение ресурсов (Dispose)."
)

add_qa_block(17,
    "Как отображаются обложки альбомов?",
    "Обложки извлекаются из ID3-тегов через TagLibSharp и отображаются в UI:\n\n"
    "1) При загрузке трека TagLibSharp читает первый фрейм изображения (APIC).\n"
    "2) Массив байт изображения декодируется в Bitmap (через System.Drawing.Common "
    "или Avalonia’s Bitmap).\n"
    "3) Bitmap отображается в Image-элементе XAML ({Binding CurrentTrack.Cover}).\n\n"
    "Важный момент — кэширование отсутствует: при переключении между треками обложка "
    "декодируется заново из файла. Для больших библиотек (1000+ треков) это может "
    "создавать задержки, поэтому в перспективе планируется добавить кэш в MemoryCache "
    "или на диск (thumbnails)."
)

add_qa_block(18,
    "Какие тесты написаны для десктопного плеера?",
    "Для десктопного плеера написано 18 unit-тестов в проекте Player.Tests (xUnit):\n"
    "• Тесты создания плейлиста и добавления треков;\n"
    "• Тесты поиска и фильтрации (регистронезависимость, частичное совпадение);\n"
    "• Тесты режимов Repeat / AutoPlay / Shuffle;\n"
    "• Тест дедупликации при добавлении одного и того же файла;\n"
    "• Тест создания PlaylistItem из файла (TagLibSharp integration test);\n"
    "• Тест сериализации/десериализации плейлиста.\n\n"
    "Не тестируются: воспроизведение (требует аудиоустройства), эквалайзер (требует "
    "аудиопотока), сканирование дисков. Это интеграционные тесты, которые планируется "
    "добавить на следующем этапе.\n\n"
    "Используется xUnit — стандартный фреймворк для .NET с параллельным запуском "
    "тестов и хорошей интеграцией с Visual Studio / Rider."
)

add_qa_block(19,
    "Какие недостатки десктопного плеера вы видите?",
    "Основные недостатки:\n"
    "1) Не чистый MVVM — ViewModel создаётся прямо в code-behind MainWindow, "
    "а не через DI-контейнер (например, Microsoft.Extensions.DependencyInjection). "
    "Для production стоит внедрить DI.\n"
    "2) Отсутствие DI — сервисы (AudioService, PlaylistService, TagReader) жёстко "
    "связаны во ViewModel. Это усложняет тестирование и замену реализаций.\n"
    "3) Нет кэширования обложек — при каждом переключении трека обложка читается "
    "с диска. Для больших библиотек это медленно.\n"
    "4) IsLoading / LoadingProgress не используются — хотя во ViewModel есть поля "
    "для отображения прогресса загрузки, они никогда не устанавливаются, оверлей "
    "загрузки не работает.\n"
    "5) Нет сохранения позиции — при закрытии плеера позиция текущего трека не "
    "сохраняется, при следующем запуске трек начинается сначала.\n"
    "6) Нет плейлистов — пользователь не может создавать несколько плейлистов "
    "(есть только один «текущий»)."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 3: Android Player
# ──────────────────────────────────────────────
add_section_header("Аудиоплеер: Android-версия (Kotlin / Jetpack Compose)", "Раздел 3")

add_qa_block(20,
    "Почему вы выбрали Jetpack Compose вместо традиционной XML-разметки?",
    "Jetpack Compose выбран по следующим причинам:\n"
    "1) Декларативный UI — описание интерфейса в виде функций (Composable), "
    "а не XML-файлов. Это сокращает количество кода и уменьшает рассогласование "
    "между логикой и представлением.\n"
    "2) Реактивность — UI автоматически перерисовывается при изменении состояния "
    "(StateFlow в ViewModel, remember/mutableState в Composable).\n"
    "3) Современность — Compose — это стандарт Android-разработки от Google, "
    "XML-разметка считается устаревшей для новых проектов.\n"
    "4) Меньше boilerplate — не нужны findViewById, ViewBinding, адаптеры "
    "для RecyclerView (вместо них LazyColumn).\n"
    "5) Совместимость с Material3 — встроенная поддержка Material Design 3 "
    "с тёмной темой и динамическими цветами (Android 12+)."
)

add_qa_block(21,
    "Как работает фоновое воспроизведение на Android?",
    "Фоновое воспроизведение реализовано через PlaybackService, наследующийся от "
    "MediaSessionService из библиотеки AndroidX Media3:\n\n"
    "1) MediaSessionService — системный сервис, который позволяет приложению "
    "продолжать воспроизведение в фоне (экран заблокирован, приложение свёрнуто).\n\n"
    "2) MediaSession — связующее звено между ExoPlayer и системой Android. "
    "Система (и другие приложения) могут управлять плеером через MediaSession — "
    "например, с экрана блокировки или из шторки уведомлений.\n\n"
    "3) ExoPlayer — ядро воспроизведения, управляется через MediaController, "
    "который ViewModel получает через MediaSession.getController().\n\n"
    "4) Кастомное уведомление — создаётся через MediaNotification.Provider "
    "с MediaStyle (кнопки Prev / PlayPause / Next). На Android 13+ обязателен "
    "вызов startForeground() с этим уведомлением, иначе система выбрасывает "
    "IllegalStateException и убивает сервис.\n\n"
    "5) Media3 обрабатывает: фокус аудио (AudioManager audio focus), "
    "прерывания от звонков/навигации, экран блокировки."
)

add_qa_block(22,
    "Какое разрешение требуется для сканирования музыки и как оно обрабатывается?",
    "Для сканирования музыкальных файлов на Android требуется разрешение "
    "READ_MEDIA_AUDIO (Android 13+, API 33) или READ_EXTERNAL_STORAGE (Android 12 и ниже).\n\n"
    "Обработка реализована в MainActivity.kt (81 строка):\n"
    "1) При запуске проверяется наличие разрешения через "
    "ContextCompat.checkSelfPermission().\n"
    "2) Если разрешения нет — запрашивается через ActivityResultLauncher "
    "(registerForActivityResult с RequestMultiplePermissions).\n"
    "3) Если пользователь согласился — запускается MusicScanner для поиска треков.\n"
    "4) Если отказался — показывается диалог с объяснением, почему разрешение "
    "необходимо, и кнопкой «Настройки» для ручного включения.\n"
    "5) На Android 13+ используется точечное разрешение READ_MEDIA_AUDIO "
    "(не READ_EXTERNAL_STORAGE), что соответствует политике конфиденциальности Google."
)

add_qa_block(23,
    "Как работает MusicScanner?",
    "MusicScanner (data/MusicScanner.kt, 68 строк) — класс для сканирования "
    "медиатеки устройства через ContentResolver и MediaStore:\n\n"
    "1) Запрос: contentResolver.query(\n"
    "    MediaStore.Audio.Media.EXTERNAL_CONTENT_URI,\n"
    "    projection = arrayOf(_ID, TITLE, ARTIST, ALBUM, DURATION, DATA),\n"
    "    selection = IS_MUSIC + \" != 0\",\n"
    "    sortOrder = TITLE + \" ASC\"\n"
    "  )\n\n"
    "2) Для каждой записи курсора создаётся объект Track с полями:\n"
    "   — id, title, artist, album, duration (в миллисекундах), uri (content://).\n\n"
    "3) Возвращается List<Track> — полный список всей музыки на устройстве.\n\n"
    "4) Выполняется в фоновом потоке (Dispatchers.IO) через suspend-функцию "
    "withContext. Результат передаётся в ViewModel через StateFlow."
)

add_qa_block(24,
    "Как устроен PlayerViewModel на Android?",
    "PlayerViewModel (ui/PlayerViewModel.kt, 310 строк) — центральный класс "
    "управления плеером. Особенности:\n\n"
    "• Наследуется от AndroidViewModel — имеет доступ к Application context.\n"
    "• Состояние UI — MutableStateFlow для каждого аспекта:\n"
    "    — tracks: List<Track> — все отсканированные треки;\n"
    "    — currentTrackIndex: Int — индекс текущего трека;\n"
    "    — isPlaying: Boolean — состояние воспроизведения;\n"
    "    — currentPosition: Long — позиция в миллисекундах;\n"
    "    — searchQuery: String — строка поиска;\n"
    "    — isShuffle, isRepeat, isAutoPlay: Boolean — режимы.\n\n"
    "• MediaController — взаимодействие с ExoPlayer через Media3:\n"
    "    — ViewModel подписывается на MediaController.Listener;\n"
    "    — при получении команд (play, pause, next, prev) обновляет StateFlow.\n\n"
    "• Управление плейлистом:\n"
    "    — setTracks(List<Track>) — загрузка треков в ExoPlayer;\n"
    "    — playTrack(index) — установка пайплайна и запуск;\n"
    "    — toggleShuffle / toggleRepeat / toggleAutoPlay — режимы.\n\n"
    "• Поиск: filterTracks(query) — фильтрация по названию/исполнителю."
)

add_qa_block(25,
    "Как реализован интерфейс плеера (PlayerScreen.kt)?",
    "PlayerScreen.kt (528 строк) — главный экран приложения, написанный "
    "полностью на Jetpack Compose. Основные компоненты:\n\n"
    "1) Верхняя панель — отображение текущего трека: название, исполнитель, "
    "обложка (или заглушка), анимированная обложка с поворотом.\n\n"
    "2) Slider позиции — ProgressIndicator с текущей позицией и длительностью. "
    "Поддержка перемотки касанием/перетаскиванием.\n\n"
    "3) Кнопки управления — Prev, Play/Pause, Next (крупные, Material3).\n\n"
    "4) Панель режимов — Repeat, AutoPlay, Shuffle (иконки с подсветкой).\n\n"
    "5) Поле поиска — TextField с фильтрацией плейлиста в реальном времени.\n\n"
    "6) LazyColumn — список всех треков с автоматической прокруткой "
    "к текущему треку (animateScrollToItem).\n\n"
    "7) BottomBar — громкость (Slider) и настройки.\n\n"
    "Вся разметка обёрнута в Scaffold с Material3 и тёмной темой "
    "(CyberBg = #0A0A0F, CyberAccent = #3A86FF)."
)

add_qa_block(26,
    "Как осуществляется навигация и автоматическая прокрутка к текущему треку?",
    "Навигация в Android-плеере однокрановая — всё приложение представляет "
    "собой один экран (Single-Activity, Single-Composable). Это упрощает "
    "архитектуру, так как плеер не требует множества экранов.\n\n"
    "Автоматическая прокрутка к текущему треку реализована через "
    "LazyListState.animateScrollToItem():\n"
    "• При изменении currentTrackIndex (StateFlow) вызывается "
    "snapshotFlow { viewModel.currentTrackIndex } в LaunchedEffect.\n"
    "• Если индекс изменился, listState.animateScrollToItem(index) "
    "плавно прокручивает список к текущему треку.\n"
    "• Это особенно удобно при включении Shuffle — текущий трек "
    "может быть в середине списка, и его нужно показать пользователю."
)

add_qa_block(27,
    "Какие есть проблемы или недостатки в Android-версии?",
    "Критические проблемы, обнаруженные при анализе:\n"
    "1) Отсутствие startForeground() в PlaybackService — на Android 13+ "
    "это вызывает IllegalStateException и краш приложения. Это самый "
    "серьёзный баг, требующий немедленного исправления.\n\n"
    "2) Плейлист полностью пересоздаётся при каждом вызове playTrack() — "
    "вместо добавления одного трека в ExoPlayer, очищается весь список "
    "и заново добавляются все отфильтрованные треки. Это неэффективно "
    "и может вызывать задержки при большом количестве треков.\n\n"
    "3) Нет сохранения последней позиции и плейлиста между сессиями — "
    "при перезапуске приложения плейлист пуст, нужно заново сканировать.\n\n"
    "4) Java 1.8 target — в build.gradle указана устаревшая версия "
    "Java, что ограничивает возможности Kotlin stdlib.\n\n"
    "5) Все разрешения запрашиваются за один раз — по правилам Google "
    "рекомендуется запрашивать разрешения по мере необходимости."
)

add_qa_block(28,
    "Чем Android-плеер отличается от десктопного?",
    "Сравнение десктопной и Android-версии плеера:\n\n"
    "• Язык: C# (.NET 10) vs Kotlin;\n"
    "• UI: Avalonia (XAML, Fluent Design) vs Jetpack Compose (декларативный, Material3);\n"
    "• Архитектура: MVVM (CommunityToolkit.Mvvm) vs MVVM (AndroidViewModel + StateFlow);\n"
    "• Аудиодвижок: NAudio (WaveOutEvent + AudioFileReader) vs Media3 (ExoPlayer);\n"
    "• Эквалайзер: Есть (BiQuadFilter, 80 Гц + 6000 Гц) vs Нет;\n"
    "• Сканирование: Файловая система + ID3-теги vs MediaStore ContentProvider;\n"
    "• Фоновое воспроизведение: Нет (только активное окно) vs Да (MediaSessionService);\n"
    "• Сохранение плейлиста: JSON в %APPDATA% vs Нет;\n"
    "• Обложки: TagLibSharp vs Нет (только текст);\n"
    "• Тесты: xUnit (18 тестов) vs Нет;\n"
    "• Количество строк: ~1200 vs ~1100.\n\n"
    "Общее: одинаковая цветовая палитра (#0A0A0F / #3A86FF), общая логика "
    "режимов воспроизведения, поиска и управления плейлистом."
)

add_qa_block(29,
    "Планируете ли вы добавлять новые функции в плеер?",
    "Да, основные направления развития:\n"
    "1) Эквалайзер на Android — перенос DSP-фильтрации из десктопной версии;\n"
    "2) Кэширование обложек — на диск или в MemoryCache для ускорения;\n"
    "3) Сохранение состояния — плейлиста и позиции между сессиями (SharedPreferences / Room);\n"
    "4) Несколько плейлистов — создание, редактирование, удаление плейлистов;\n"
    "5) Синхронизация — облачная синхронизация плейлистов между устройствами;\n"
    "6) Радио / подкасты — интеграция с потоковыми сервисами;\n"
    "7) EqualizerView — графическое отображение спектра;\n"
    "8) GestureControl — свайпы для переключения треков, изменение громкости "
    "по краю экрана (как в VLC)."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 4: Backend
# ──────────────────────────────────────────────
add_section_header("Интернет-магазин: Бэкенд (Django REST)", "Раздел 4")

add_qa_block(30,
    "Почему вы выбрали Django для бэкенда, а не FastAPI или Flask?",
    "Django выбран по нескольким причинам:\n"
    "1) «Всё включено» — Django предоставляет ORM, админ-панель, "
    "аутентификацию, сериализацию, миграции и интернационализацию «из коробки». "
    "Для Flask или FastAPI пришлось бы подключать SQLAlchemy, Alembic, "
    "Flask-Admin, Flask-Login и т.д.\n\n"
    "2) DRF (Django REST Framework) — де-факто стандарт для REST API на Python. "
    "Предоставляет ViewSet, Serializer, Permission, Pagination и Browsable API.\n\n"
    "3) Админ-панель — Django Admin сокращает время разработки админки на ~80%, "
    "а кастомная админка /cp/ для специфических задач не требует отдельного UI-фреймворка.\n\n"
    "4) ORM — декларативное описание моделей с managed=False для legacy-таблиц. "
    "Django ORM поддерживает сложные запросы с select_related, annotate, F().\n\n"
    "5) Сообщество и документация — Django существует с 2005 года, огромное "
    "количество пакетов (simplejwt, corsheaders, debug-toolbar, silk).\n\n"
    "FastAPI рассматривался, но не подошёл из-за отсутствия встроенной "
    "админ-панели и более молодого сообщества."
)

add_qa_block(31,
    "Опишите структуру эндпоинтов REST API.",
    "API имеет 9 эндпоинтов, сгруппированных по функционалу:\n\n"
    "Товары и артисты (публичные):\n"
    "• GET /api/artists/ — список всех артистов;\n"
    "• GET /api/artists/<slug>/ — детальная информация об артисте + его товары;\n"
    "• GET /api/merch/ — список всех товаров (с фильтром ?artist_slug=...);\n\n"
    "Аутентификация:\n"
    "• POST /api/auth/register/ — регистрация нового пользователя;\n"
    "• POST /api/auth/login/ — вход, возвращает JWT access + refresh;\n"
    "• POST /api/auth/token/refresh/ — обновление access-токена;\n"
    "• GET/PUT /api/auth/profile/ — получение/обновление профиля;\n\n"
    "Заказы (требуют аутентификации):\n"
    "• POST /api/checkout/ — оформление заказа (корзина → Order);\n\n"
    "Все эндпоинты используют APIView (class-based views). "
    "Права доступа: AllowAny для публичных эндпоинтов, IsAuthenticated "
    "для checkout и профиля."
)

add_qa_block(32,
    "Как работает JWT-аутентификация?",
    "JWT-аутентификация реализована через djangorestframework-simplejwt:\n\n"
    "1) При регистрации (POST /api/auth/register/) создаётся пользователь "
    "ShopUser и UserProfile, затем генерируется пара токенов "
    "(access + refresh) через RefreshToken.for_user(user).\n\n"
    "2) При входе (POST /api/auth/login/) проверяется username и password, "
    "при успехе — возвращаются access и refresh токены.\n\n"
    "3) Access-токен живёт 2 часа (ACCESS_TOKEN_LIFETIME=timedelta(hours=2)). "
    "Используется для аутентификации запросов в заголовке: "
    "Authorization: Bearer <token>.\n\n"
    "4) Refresh-токен живёт 7 дней (REFRESH_TOKEN_LIFETIME=timedelta(days=7)). "
    "Передаётся в POST /api/auth/token/refresh/ для получения нового access-токена "
    "без повторного ввода пароля.\n\n"
    "5) На фронтенде токены хранятся в localStorage. AuthContext перехватывает "
    "401-ответы и автоматически обновляет access-токен через refresh.\n\n"
    "6) В настройках DRF: DEFAULT_AUTHENTICATION_CLASSES установлен "
    "на JWTAuthentication, но DEFAULT_PERMISSION_CLASSES — AllowAny "
    "(права настраиваются индивидуально для каждого эндпоинта)."
)

add_qa_block(33,
    "Расскажите о моделях базы данных и их связях.",
    "Всего 10 моделей в приложении shop:\n\n"
    "Unmanaged (managed=False, внешние таблицы, Django НЕ управляет миграциями):\n"
    "• ShopUser (db_table='users_user') — пользователи. НЕ наследуется от "
    "AbstractUser, все поля (username, password, email, role, balance и др.) "
    "объявлены вручную.\n"
    "• Creator (shop_creator) — исполнители (10 артистов).\n"
    "• Category (shop_category) — категории товаров с self-referencing parent_id.\n"
    "• Product (shop_product) — товары. FK на Creator и Category.\n"
    "• ProductVariant (shop_productvariant) — варианты (размер + сток). FK на Product.\n"
    "• ProductImage (shop_productimage) — изображения. FK на Product.\n\n"
    "Managed (managed=True, Django управляет миграциями):\n"
    "• UserProfile (shop_userprofile) — профиль пользователя (связь через IntegerField user_id).\n"
    "• Order (shop_order) — заказы. user_id — IntegerField (не FK, так как ShopUser — unmanaged).\n"
    "• OrderItem (shop_orderitem) — позиции заказа. FK на Order. CASCADE при удалении.\n"
    "• Supply (shop_supply) — привозы (поступления товара). FK на ProductVariant. CASCADE.\n\n"
    "Связи: Creator 1→N Product, Category 1→N Product, Product 1→N ProductVariant, "
    "Product 1→N ProductImage, Order 1→N OrderItem, ProductVariant 1→N Supply."
)

add_qa_block(34,
    "Что означает managed=False и зачем вы его используете?",
    "managed=False — это атрибут Meta-класса модели Django, который указывает, "
    "что Django не должен управлять жизненным циклом этой таблицы в базе данных:\n"
    "• Не создавать таблицу при миграции (makemigrations/migrate);\n"
    "• Не удалять таблицу;\n"
    "• Не изменять схему таблицы;\n\n"
    "Мы используем managed=False для 6 моделей (ShopUser, Creator, Category, "
    "Product, ProductVariant, ProductImage), потому что:\n"
    "• Таблицы были созданы внешней системой (или более старой версией БД);\n"
    "• Мы не хотим, чтобы Django случайно изменил их структуру;\n"
    "• Django ORM всё равно может читать и писать данные — managed влияет "
    "только на управление схемой.\n\n"
    "При миграции с PostgreSQL на SQLite мы временно установили managed=True "
    "для этих моделей, выполнили миграцию (создали таблицы), перенесли данные, "
    "и вернули managed=False обратно."
)

add_qa_block(35,
    "Как работает оформление заказа (Checkout)?",
    "Оформление заказа реализовано в CheckoutView (APIView, только IsAuthenticated):\n\n"
    "1) Клиент отправляет POST /api/checkout/ с JSON:\n"
    "   { \"items\": [{ \"id\": 1, \"quantity\": 2 }, ...] }\n\n"
    "2) Сервер проходит по каждому товару:\n"
    "   — Ищет ProductVariant по product_id (через variants.first());\n"
    "   — Проверяет достаточно ли stock (variant.stock >= quantity);\n"
    "   — Уменьшает variant.stock на quantity (через обычное присваивание, не F());\n"
    "   — Добавляет позицию в список order_items.\n\n"
    "3) Создаёт Order (user_id=request.user.id, total=..., status='processing').\n\n"
    "4) Для каждой позиции создаёт OrderItem.\n\n"
    "5) Возвращает JSON с order_id, сообщением об успехе.\n\n"
    "Проблемы текущей реализации:\n"
    "• Используется variants.first() — учитывается только первый вариант товара;\n"
    "• Нет транзакционной обёртки — при сбое между уменьшением stock и созданием "
    "заказа данные могут быть повреждены;\n"
    "• Нет F() для атомарного обновления stock — возможен race condition "
    "при параллельных запросах."
)

add_qa_block(36,
    "Зачем нужны две админ-панели — /admin/ и /cp/?",
    "В проекте реализованы две административные панели для разных целей:\n\n"
    "1) Стандартная Django Admin (/admin/):\n"
    "   — Автоматически генерируется на основе моделей (регистрация через @admin.register);\n"
    "   — Даёт CRUD для всех 10 моделей (ShopUser, Creator, Product, Order и т.д.);\n"
    "   — Требует минимальных настроек (list_display, list_filter, search_fields);\n"
    "   — Недостаток: стандартный дизайн, нет кастомных действий (дашборд, статистика).\n\n"
    "2) Кастомная админ-панель (/cp/):\n"
    "   — 16 function-based views, оформленных в едином тёмном стиле;\n"
    "   — Дашборд (dashboard) — сводка: количество заказов, выручка, товары, "
    "   последние заказы, товары с низким остатком;\n"
    "   — CRUD для артистов, товаров, категорий, привозов, заказов, пользователей;\n"
    "   — Изменение статуса заказа (processing → shipped → delivered);\n"
    "   — Привозы: добавление поступления товара с атомарным обновлением "
    "   variant.stock через F('stock') + quantity;\n"
    "   — Тёмная тема (12 HTML-шаблонов).\n\n"
    "Разделение позволяет: использовать стандартную админку для быстрого доступа "
    "и кастомную для ежедневной работы сотрудников магазина."
)

add_qa_block(37,
    "Как реализован функционал привозов (Supply)?",
    "Привозы (Supply) — это модель, отслеживающая поступление товаров на склад. "
    "Реализация:\n\n"
    "Модель Supply (shop_supply):\n"
    "• variant (FK → ProductVariant, CASCADE) — какой вариант товара поступил;\n"
    "• quantity — количество поступивших единиц;\n"
    "• supplier — поставщик;\n"
    "• purchase_price — закупочная цена (опционально);\n"
    "• notes — примечания;\n"
    "• created_by — кто создал запись.\n\n"
    "Логика привоза в SupplyList (admin_views.py):\n"
    "1) Выбирается вариант товара, указывается количество;\n"
    "2) variant.stock = F('stock') + quantity — атомарное обновление "
    "(F() expression гарантирует, что не будет race condition при параллельных запросах);\n"
    "3) variant.save() — сохраняет новое количество;\n"
    "4) Supply.objects.create(...) — создаёт запись о привозе.\n\n"
    "Благодаря CASCADE при удалении варианта товара все связанные привозы "
    "удаляются автоматически."
)

add_qa_block(38,
    "Какие сериализаторы используются и как они работают?",
    "В проекте 9 сериализаторов (serializers.py, 170 строк):\n\n"
    "ModelSerializer:\n"
    "• CreatorListSerializer — базовые поля артиста + avatar_color (вычисляется) "
    "+ is_musician (всегда true);\n"
    "• CreatorDetailSerializer — артист + его товары (через SerializerMethodField items);\n"
    "• ProductItemSerializer — товар с вычисляемыми полями: artist, size, stock, "
    "sticker_color (по маппингу slug → цвет), condition, image_sticker_type;\n"
    "• OrderItemSerializer — позиция заказа: product_name, price, quantity;\n"
    "• OrderSerializer — заказ + вложенные items.\n\n"
    "Serializer (не ModelSerializer):\n"
    "• RegisterSerializer — валидация username (уникальность), email, password, "
    "display_name;\n"
    "• LoginSerializer — username + password;\n"
    "• ProfileUpdateSerializer — display_name, favorite_artist;\n"
    "• CheckoutSerializer — список items (id + quantity), каждый валидируется "
    "через CheckoutItemSerializer (id — IntegerField, quantity — min 1).\n\n"
    "Маппинг цветов: COLOR_MAP — словарь slug → цвет (gonefludd → 'pink', "
    "lsp → 'cyan', pharaoh → 'orange' и т.д.). Всего 10 цветов для 10 артистов "
    "(плюс FALLBACK_COLORS для неизвестных)."
)

add_qa_block(39,
    "Какие конфигурационные настройки Django вы кастомизировали?",
    "Основные кастомизации в settings.py:\n\n"
    "• INSTALLED_APPS — добавлены: corsheaders, rest_framework, "
    "rest_framework_simplejwt, shop (приложение).\n\n"
    "• MIDDLEWARE — добавлен CorsMiddleware (для CORS-заголовков).\n\n"
    "• REST_FRAMEWORK — DEFAULT_AUTHENTICATION_CLASSES = JWTAuthentication, "
    "DEFAULT_PERMISSION_CLASSES = AllowAny.\n\n"
    "• SIMPLE_JWT — access token: 2 часа, refresh token: 7 дней.\n\n"
    "• CORS_ALLOW_ALL_ORIGINS = True — разрешены запросы с любых доменов "
    "(для разработки). В production нужно ограничить конкретными доменами.\n\n"
    "• CSRF_TRUSTED_ORIGINS — список доверенных origins для CSRF-защиты.\n\n"
    "• DATABASES — переключено с PostgreSQL на SQLite (django.db.backends.sqlite3).\n\n"
    "• DEBUG = True — включен режим отладки (только для разработки).\n\n"
    "• ALLOWED_HOSTS = ['*'] — разрешены все хосты (только для разработки)."
)

add_qa_block(40,
    "Как обеспечивается безопасность API?",
    "Текущий уровень безопасности:\n\n"
    "Реализовано:\n"
    "• JWT-аутентификация — access-токены с ограниченным сроком (2 часа), "
    "refresh-токены для продления сессии;\n"
    "• Проверка прав — IsAuthenticated для checkout и профиля;\n"
    "• Хеширование паролей — Django make_password() (PBKDF2 + SHA-256);\n"
    "• CSRF-защита — включена для session-based endpoints (admin);\n"
    "• CORS — ограничен список доверенных origins для CSRF.\n\n"
    "Проблемы (только для разработки):\n"
    "• DEBUG = True — показывает traceback и конфигурацию при ошибках;\n"
    "• ALLOWED_HOSTS = ['*'] — уязвимость к Host header attacks;\n"
    "• CORS_ALLOW_ALL_ORIGINS = True — любой сайт может делать запросы к API;\n"
    "• SECRET_KEY в открытом виде — в settings.py (должен быть в .env);\n"
    "• Пароль БД был в открытом виде (до миграции на SQLite);\n"
    "• Нет rate limiting — возможна брутфорс-атака на /auth/login/;\n"
    "• Нет HTTPS — все данные передаются в открытом виде.\n\n"
    "Для production обязательно: DEBUG=False, ALLOWED_HOSTS — конкретные домены, "
    "SECRET_KEY в переменных окружения, CORS_ORIGIN_WHITELIST, HTTPS, rate limiting."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 5: Frontend
# ──────────────────────────────────────────────
add_section_header("Интернет-магазин: Веб-фронтенд (React / TypeScript)", "Раздел 5")

add_qa_block(41,
    "Почему React, а не Vue или Angular?",
    "React выбран по следующим причинам:\n"
    "1) Популярность — React — самая используемая библиотека для SPA, "
    "огромное сообщество, множество готовых решений и компонентов.\n"
    "2) TypeScript — React имеет отличную поддержку TypeScript, что "
    "обеспечивает статическую типизацию, автодополнение и раннее обнаружение ошибок.\n"
    "3) Vite — современный сборщик с быстрым HMR (Hot Module Replacement) "
    "и оптимизированной production-сборкой. Vite 5.2.11 даёт значительный "
    "прирост скорости по сравнению с Create React App (Webpack).\n"
    "4) React Router v7 — последняя версия с поддержкой loaders, actions "
    "и вложенных маршрутов.\n"
    "5) Shared code с React Native — можно переиспользовать типы, утилиты "
    "и хуки между вебом и мобильным приложением.\n\n"
    "Vue и Angular не рассматривались, так как проект начат на React, "
    "и React Native является естественным расширением для мобильной версии."
)

add_qa_block(42,
    "Как устроена маршрутизация в приложении?",
    "Маршрутизация реализована через React Router v7 (9 маршрутов):\n\n"
    "• / — Главная страница (Hero-баннер, бегущая строка артистов, "
    "блоки с товарами, акции);\n"
    "• /catalog — Каталог (фильтрация по артисту, поиск, сортировка);\n"
    "• /catalog/:slug — Детальная страница товара (описание, размеры, "
    "сток, кнопка «В корзину»);\n"
    "• /cart — Корзина (список товаров, количество, итоговая сумма, "
    "оформление заказа);\n"
    "• /wishlist — Избранное (сетка товаров, удаление, покупка);\n"
    "• /about — О нас (информация о проекте);\n"
    "• /profile — Профиль (статистика, история заказов);\n"
    "• /login — Вход;\n"
    "• /register — Регистрация.\n\n"
    "Маршруты организованы через createBrowserRouter в App.tsx. "
    "Header отображается на всех страницах (Layout route)."
)

add_qa_block(43,
    "Как работает корзина и избранное?",
    "Корзина (cart) и избранное (wishlist) реализованы поверх localStorage с "
    "состоянием в App.tsx (единый компонент-контейнер):\n\n"
    "Корзина:\n"
    "• Хранится как CartItem[] — каждый элемент: { id, name, price, quantity, image? };\n"
    "• Функции: addToCart, removeFromCart, updateQuantity, clearCart;\n"
    "• При изменении корзина сохраняется в localStorage;\n"
    "• На странице /cart отображается список с итоговой суммой (total = sum price*quantity);\n"
    "• При оформлении заказа (POST /api/checkout/) отправляется только список id + quantity;\n"
    "• После успешного заказа корзина очищается.\n\n"
    "Избранное:\n"
    "• Хранится как Set<number> (id товаров);\n"
    "• Функции: toggleWishlist, isInWishlist;\n"
    "• Отображается на /wishlist в виде сетки товаров;\n"
    "• Из избранного можно добавить товар в корзину или удалить.\n\n"
    "Оба механизма работают без серверной части — они полностью клиентские."
)

add_qa_block(44,
    "Как организована дизайн-система и стилизация?",
    "Дизайн-система организована в едином файле index.css (1688 строк):\n\n"
    "• CSS-переменные (custom properties) в :root — полная тёмная тема:\n"
    "   --bg-primary: #0A0A0F (фоновый),\n"
    "   --bg-secondary: #1A1A2E (карточки),\n"
    "   --accent: #3A86FF (акцентный синий),\n"
    "   --text-primary: #FFFFFF,\n"
    "   --text-secondary: #B0B0B0 и т.д.\n\n"
    "• Компонентный подход — стили сгруппированы по компонентам: header, "
    "hero, catalog, card, cart, forms, buttons, modals и т.д.\n\n"
    "• Анимации — @keyframes для бегущей строки (marquee), появления "
    "карточек (fadeIn, slideUp), подсветки кнопок.\n\n"
    "• Адаптивность — минимальная, сайт ориентирован на desktop. "
    "Media queries есть только для крупных планшетов.\n\n"
    "• SVG-иконки — 53 кастомные SVG-иконки + 10 иконок артистов, "
    "встроенные в CustomSvg.tsx и ArtistIcons.tsx.\n\n"
    "Плюсы: всё в одном файле, легко найти. Минусы: нет модульности "
    "(CSS Modules / Tailwind), 1688 строк сложно поддерживать."
)

add_qa_block(45,
    "Как реализована интеграция с бэкендом (API-клиент)?",
    "API-клиент реализован в App.tsx и AuthContext.tsx с использованием "
    "нативного fetch() без дополнительных библиотек (Axios не используется):\n\n"
    "• Базовый URL: http://127.0.0.1:8001/api/ (внимание: порт 8001, "
    "хотя бэкенд по умолчанию на 8000 — это известная нестыковка).\n\n"
    "• Загрузка данных (артисты, товары) — при монтировании App.tsx "
    "вызываются fetch() и результаты сохраняются в state.\n\n"
    "• Fallback на мок-данные — если API недоступен (fetch не удался), "
    "загружаются моки из констант. Это позволяет разрабатывать UI без запущенного бэкенда.\n\n"
    "• Аутентификация — AuthContext перехватывает fetch и добавляет "
    "Authorization: Bearer <token> из localStorage. При 401 вызывает "
    "refresh-эндпоинт и повторяет запрос.\n\n"
    "• Состояние загрузки — isLoaded / loading флаги для отображения "
    "спиннера во время fetch."
)

add_qa_block(46,
    "Как работает AuthContext?",
    "AuthContext (context/AuthContext.tsx, 176 строк) — центральный "
    "компонент для управления аутентификацией:\n\n"
    "• Хранит: user (объект пользователя), accessToken, refreshToken, "
    "isAuthenticated, isLoading.\n\n"
    "• login(username, password) — POST /api/auth/login/ → сохраняет "
    "токены и пользователя в state + localStorage.\n\n"
    "• register(username, email, password, displayName) — "
    "POST /api/auth/register/ → автоматически логинит после регистрации.\n\n"
    "• logout() — очищает state и localStorage (но не инвалидирует refresh-токен "
    "на сервере, это известное ограничение simplejwt).\n\n"
    "• Auto-refresh — при любом 401-ответе от API, AuthContext перехватывает "
    "ошибку, вызывает /api/auth/token/refresh/ и обновляет access-токен. "
    "Если refresh не удался — принудительный logout.\n\n"
    "• Provider — оборачивает всё приложение, предоставляя контекст "
    "через React.createContext."
)

add_qa_block(47,
    "Как реализованы мок-данные?",
    "Мок-данные — это константы с полной структурой артистов и товаров, "
    "используемые как fallback при недоступности API. Они содержат:\n\n"
    "• 10 артистов с теми же slug (lsp, gonefludd, pharaoh, oxxxymiron и т.д.);\n"
    "• 66 товаров (чуть больше, чем в seed_pg.py с 62 товарами, с некоторыми "
    "расхождениями в ценах и количестве);\n"
    "• Цвета артистов (аналог COLOR_MAP на бэкенде);\n"
    "• Полноценные описания и характеристики.\n\n"
    "Мок-данные позволили разрабатывать и тестировать UI без запуска "
    "бэкенда и базы данных. Однако они дублируются с реальными данными "
    "и могут расходиться при изменении seed-скрипта."
)

add_qa_block(48,
    "Какие страницы реализованы в веб-приложении?",
    "Реализовано 9 страниц (pages/):\n\n"
    "1) Home.tsx — главная: Hero-баннер с киберпанк-стилем, бегущая строка "
    "с именами артистов, блоки «Новинки», «Хиты продаж», «Специальные "
    "предложения» с анимированными карточками.\n\n"
    "2) Catalog.tsx — каталог: сетка товаров (CSS Grid), фильтр по артисту "
    "(dropdown), поиск по названию, сортировка (по цене, по дате).\n\n"
    "3) ProductDetail.tsx — детальная карточка: название, цена, описание, "
    "выбор размера, остаток на складе, кнопка «В корзину», избранное.\n\n"
    "4) Cart.tsx — корзина: таблица товаров, количество (плюс/минус), "
    "удаление, итоговая сумма, кнопка «Оформить заказ».\n\n"
    "5) Wishlist.tsx — избранное: сетка товаров с сердечком, удаление, "
    "добавление в корзину.\n\n"
    "6) Profile.tsx — профиль: данные пользователя, статистика (количество "
    "заказов, потрачено), история заказов.\n\n"
    "7) Login.tsx / Register.tsx — формы входа и регистрации.\n\n"
    "8) About.tsx — информация о проекте."
)

add_qa_block(49,
    "Какие проблемы есть во фронтенде?",
    "Основные проблемы:\n"
    "1) Жёстко зашит порт 8001 — в App.tsx и AuthContext.tsx указан "
    "http://127.0.0.1:8001, но бэкенд по умолчанию запускается на 8000. "
    "Это несовместимость, требующая исправления (либо менять порт бэкенда, "
    "либо URL во фронтенде).\n\n"
    "2) Один CSS-файл на 1688 строк — отсутствует модульность. "
    "Рекомендуется разбить на CSS Modules или использовать Tailwind.\n\n"
    "3) Нет адаптивной вёрстки — мобильные устройства отображаются "
    "неработоспособно (горизонтальная прокрутка, мелкие элементы).\n\n"
    "4) Мок-данные расходятся с реальными — 66 товаров в моках vs 62 "
    "в seed-скрипте, отличаются цены.\n\n"
    "5) Неиспользуемые зависимости — framer-motion и lucide-react "
    "установлены, но не используются, увеличивая размер бандла.\n\n"
    "6) Нет серверного рендеринга (SSR) — SEO-оптимизация отсутствует, "
    "поисковые системы не видят контент страниц (пустой HTML до загрузки JS)."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 6: React Native
# ──────────────────────────────────────────────
add_section_header("Интернет-магазин: Мобильное приложение (React Native)", "Раздел 6")

add_qa_block(50,
    "Зачем React Native, если уже есть Android-плеер?",
    "React Native (MerchMarketRN) — это мобильный клиент интернет-магазина "
    "(стикеры и мерч), а Android-плеер — это отдельное приложение для "
    "воспроизведения музыки. Они решают разные задачи:\n"
    "• Android-плеер (Kotlin/Compose): воспроизведение локальной музыки, "
    "сканирование MediaStore, эквалайзер, фоновое воспроизведение.\n"
    "• React Native (TypeScript/Expo): покупка мерча, просмотр каталога, "
    "корзина, оформление заказов, профиль.\n\n"
    "React Native позволяет переиспользовать до 60-70% кода с веб-фронтендом "
    "(типы, логика, утилиты), что значительно ускоряет разработку по сравнению "
    "с написанием нативного Android-приложения для магазина."
)

add_qa_block(51,
    "Почему вы не использовали React Navigation?",
    "В MerchMarketRN реализована кастомная стековая навигация (pushView/goBack) "
    "вместо React Navigation по нескольким причинам:\n"
    "1) Учебный проект — целью было изучение устройства навигации "
    "«изнутри», а не использование готового решения.\n"
    "2) Лёгкость — для 9 экранов кастомная навигация проще и не требует "
    "дополнительных зависимостей.\n"
    "3) Полный контроль — анимации переходов, жест «назад», "
    "передача параметров между экранами.\n\n"
    "Недостатки: нет поддержки глубоких ссылок (deep linking), "
    "нет tab-навигации, нет header с автоматической интеграцией "
    "с системной кнопкой «Назад» на Android."
)

add_qa_block(52,
    "Как реализовано боковое меню (drawer)?",
    "Боковое меню реализовано с помощью react-native-reanimated "
    "и Animated API:\n\n"
    "1) Drawer — нативный компонент, использующий useAnimatedStyle "
    "для анимации translateX при открытии/закрытии.\n\n"
    "2) Trigger — иконка «гамбургер» в Header, при нажатии "
    "вызывает setIsDrawerOpen(true).\n\n"
    "3) Затемнение — overlay с opacity-анимацией на фоне.\n\n"
    "4) Пункты меню: Главная, Каталог, Корзина, Избранное, О нас, "
    "Профиль, Войти/Выйти.\n\n"
    "5) Закрытие — по нажатию на пункт меню, на оверлей, "
    "или свайпом влево."
)

add_qa_block(53,
    "Как React Native приложение получает данные?",
    "Данные загружаются из двух источников:\n\n"
    "1) HTTP-клиент (api/client.ts):\n"
    "   — Базовый URL: http://10.0.2.2:8001/api/ (10.0.2.2 — localhost "
    "с точки зрения Android-эмулятора);\n"
    "   — GET /artists/, /merch/ — загрузка каталога;\n"
    "   — POST /auth/login/, /auth/register/ — аутентификация;\n"
    "   — jwt-токен в AsyncStorage (аналог localStorage на React Native);\n"
    "   — Auto-refresh при 401.\n\n"
    "2) Мок-данные (fallback):\n"
    "   — Если fetch не удался (нет интернета / сервер не запущен), "
    "используются те же мок-константы, что и в веб-фронтенде;\n"
    "   — Это обеспечивает работоспособность приложения при "
    "демонстрации без запущенного бэкенда.\n\n"
    "Состояние хранится в контексте App.tsx (аналог веб-версии)."
)

add_qa_block(54,
    "Какие экраны реализованы в React Native?",
    "Реализовано 9 экранов (screens/), полностью соответствующих "
    "веб-версии:\n\n"
    "1) HomeScreen — главная: Hero с градиентом, бегущая строка "
    "(Animated + loop), карточки товаров.\n\n"
    "2) CatalogScreen — каталог: FlatList с фильтрацией по артисту.\n\n"
    "3) ProductDetailScreen — детальная карточка: изображение, описание, "
    "выбор размера, цена, кнопка «В корзину».\n\n"
    "4) CartScreen — корзина: FlatList с количеством, итог, checkout.\n\n"
    "5) WishlistScreen — избранное: GridView, добавление в корзину.\n\n"
    "6) ProfileScreen — профиль + история заказов.\n\n"
    "7) LoginScreen / RegisterScreen — формы входа/регистрации.\n\n"
    "8) AboutScreen — о проекте.\n\n"
    "Технологии: Expo SDK 56, React Native 0.85.3, "
    "react-native-svg (иконки), react-native-reanimated (анимации), "
    "expo-linear-gradient (фоны)."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 7: Database & Infrastructure
# ──────────────────────────────────────────────
add_section_header("База данных и инфраструктура", "Раздел 7")

add_qa_block(55,
    "Почему вы перенесли БД с PostgreSQL на SQLite и как это делали?",
    "Перенос с PostgreSQL на SQLite выполнен для упрощения развёртывания "
    "проекта:\n\n"
    "• PostgreSQL требует установки отдельного сервера базы данных, "
    "настройки пользователей и прав доступа;\n"
    "• SQLite — файловая БД, не требует сервера, всё хранится в одном "
    "файле (db.sqlite3);\n"
    "• Для учебного/демонстрационного проекта SQLite более чем достаточен "
    "(нет параллельного доступа, объём данных маленький — 243 записи).\n\n"
    "Процесс миграции:\n"
    "1) Дамп PostgreSQL: pg_dump -U postgres -d merchmarket_db -X utf8 > dump.sql\n"
    "2) Парсинг дампа — извлечение INSERT-запросов для всех 10 таблиц;\n"
    "3) Изменение ENGINE в settings.py с postgresql на sqlite3;\n"
    "4) Временное переключение 6 unmanaged-моделей на managed=True;\n"
    "5) python manage.py makemigrations → 0001_initial.py;\n"
    "6) python manage.py migrate — создание таблиц;\n"
    "7) Возврат managed=False;\n"
    "8) Загрузка данных из дампа (INSERT) — 243 объекта;\n"
    "9) Создание seed.py на основе seed_pg.py;\n"
    "10) Проверка: system check без ошибок, API работает."
)

add_qa_block(56,
    "Сколько данных в базе и какие таблицы?",
    "База данных SQLite (db.sqlite3) содержит 243 объекта в 10 таблицах:\n\n"
    "Unmanaged (6 таблиц, 200+ записей):\n"
    "• users_user (ShopUser): 3 пользователя (admin, testuser, testuser2);\n"
    "• shop_creator (Creator): 10 артистов (ЛСП, GONE.Fludd, Pharaoh, "
    "Oxxxymiron, Скриптонит, Гуф, Miyagi, Boulevard Depo, ATL, FACE);\n"
    "• shop_category (Category): 7 категорий (Футболки, Штаны, Худи, "
    "Кепки, Винил, Кассеты, Аксессуары);\n"
    "• shop_product (Product): 77 товаров (официальные + секонд-хенд);\n"
    "• shop_productvariant (ProductVariant): 143 варианта (размер/сток);\n"
    "• shop_productimage (ProductImage): без данных (таблица пуста, "
    "изображения задаются через slug артиста).\n\n"
    "Managed (4 таблицы, ~40 записей):\n"
    "• shop_userprofile (UserProfile): 3 профиля пользователей;\n"
    "• shop_order (Order): 0 заказов (после миграции);\n"
    "• shop_orderitem (OrderItem): 0 позиций;\n"
    "• shop_supply (Supply): 0 привозов.\n\n"
    "django_migrations: 1 миграция (0001_initial).\n"
    "django_content_type, auth_permission и другие системные таблицы Django."
)

add_qa_block(57,
    "Как работает seed-скрипт?",
    "seed.py (shop/management/commands/seed.py, 241 строка) — "
    "команда Django management для наполнения базы данными:\n\n"
    "Особенности:\n"
    "• Идемпотентность — проверяет существование данных перед вставкой "
    "(по slug для артистов, по name+creator_id для товаров). "
    "Можно запускать многократно без дублирования.\n\n"
    "• Использование сырого SQL — через connection.cursor() вместо "
    "Django ORM. Это связано с тем, что модели с managed=False не "
    "поддерживают ORM-операции Django (save(), create()) в полном объёме.\n\n"
    "• Категории — создаются автоматически (7 категорий, REQUIRED_CATEGORIES).\n\n"
    "• Артисты — 10 артистов с полными данными: описание, slug, "
    "путь к логотипу (/media/logos/{slug}.png).\n\n"
    "• Товары — 77 товаров с вариантами (размер, сток, SKU).\n\n"
    "• Форматы цен — Decimal для точности, даты — timezone-aware datetime.\n\n"
    "• RETURNING — используется PostgreSQL-синтаксис (RETURNING id для "
    "получения ID созданной записи). Для SQLite 3.45.1+ это тоже работает."
)

add_qa_block(58,
    "Как организован запуск проекта?",
    "Для упрощения запуска разработаны BAT-файлы в корне проекта:\n\n"
    "• run_backend.bat — запуск Django-сервера: активация venv, "
    "python manage.py runserver 127.0.0.1:8000;\n\n"
    "• run_frontend.bat — запуск Vite dev server: cd frontend && npm run dev;\n\n"
    "• run_mobile.bat — запуск React Native (Expo): cd MerchMarketRN && npx expo start;\n\n"
    "• start_all.bat — запуск всех трёх компонентов одновременно "
    "(открывает три консольных окна);\n\n"
    "• stop_all.bat — остановка всех процессов (taskkill по именам окон).\n\n"
    "Порты:\n"
    "• Бэкенд: 127.0.0.1:8000 (Django);\n"
    "• Фронтенд: localhost:5173 (Vite);\n"
    "• React Native: порт 8081 (Metro bundler).\n\n"
    "Примечание: во фронтенде жёстко зашит порт 8001 для API "
    "(App.tsx, AuthContext), что не соответствует порту 8000 бэкенда. "
    "Необходимо исправить либо settings.py (изменить порт), "
    "либо URL во фронтенде."
)

add_qa_block(59,
    "Почему вы выбрали BAT-файлы вместо PowerShell?",
    "BAT-файлы выбраны по причине проблем с кодировкой кириллицы "
    "в PowerShell при передаче аргументов.\n\n"
    "• Путь проекта содержит кириллицу («E:\\илья проект»);\n"
    "• PowerShell по-разному обрабатывает пути с кириллицей в зависимости "
    "от версии и кодировки консоли;\n"
    "• BAT-файлы (cmd.exe) обрабатывают кириллические пути корректно "
    "в кодировке OEM (866) на русской Windows.\n\n"
    "Все BAT-файлы работают через pushd (временная смена директории) "
    "для корректной обработки кириллического пути."
)

add_qa_block(60,
    "Какой .gitignore у проекта?",
    "В проекте есть .gitignore в корне со следующим содержимым:\n\n"
    "• Сборки: **/bin/, **/obj/, **/build/, **/.gradle/, **/local.properties;\n"
    "• Python: __pycache__/, *.pyc, venv/, *.sqlite3;\n"
    "• Node: node_modules/, *.tsbuildinfo, vite.config.ts.timestamp-*;\n"
    "• IDE: .idea/, *.user, *.suo;\n"
    "• OS: Thumbs.db, .DS_Store;\n"
    "• Logs: *.log;\n"
    "• React Native: **/Exponent.app, **/android/app/build/.\n\n"
    "Отсутствует: игнорирование .env-файлов (SECRET_KEY и пароли "
    "до сих пор в открытом виде в settings.py)."
)

add_qa_block(61,
    "Есть ли CI/CD?",
    "На данный момент CI/CD не настроен. Проект собирается и "
    "развёртывается вручную.\n\n"
    "Рекомендации по CI/CD:\n"
    "• GitHub Actions для автоматического запуска тестов при push;\n"
    "• Docker-контейнеризация (Dockerfile для бэкенда, фронтенда);\n"
    "• Автоматическая сборка APK (Android) через GitHub Actions;\n"
    "• Деплой бэкенда на Railway / Render / Vercel (Python);\n"
    "• Деплой фронтенда на Vercel / Netlify;\n"
    "• Использование GitHub Secrets для хранения SECRET_KEY и паролей.\n\n"
    "Это одно из ключевых направлений для production-подготовки."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 8: Architecture, Security & Testing
# ──────────────────────────────────────────────
add_section_header("Архитектура, безопасность и тестирование", "Раздел 8")

add_qa_block(62,
    "Какие архитектурные паттерны вы использовали?",
    "Паттерны, использованные в проекте:\n\n"
    "1) MVVM (Model-View-ViewModel):\n"
    "   — Desktop: CommunityToolkit.Mvvm с source-генераторами;\n"
    "   — Android: AndroidViewModel + StateFlow + Compose.\n"
    "   Разделение UI, бизнес-логики и данных. ViewModel не зависит "
    "от View и может быть протестирована.\n\n"
    "2) MVC (Model-View-Controller) — Django:\n"
    "   — Model: ORM (10 моделей);\n"
    "   — View: APIView (бизнес-логика);\n"
    "   — Controller: URL router + DRF Serializer.\n\n"
    "3) Service Layer — Android MediaSessionService:\n"
    "   — Фоновый сервис для воспроизведения, отделённый от UI.\n\n"
    "4) Repository — MusicScanner:\n"
    "   — Абстракция над ContentResolver/MediaStore, предоставляет "
    "чистый интерфейс для получения List<Track>.\n\n"
    "5) Bridge / Adapter — managed=False модели в Django:\n"
    "   — Адаптация внешней схемы БД под Django ORM без изменения таблиц.\n\n"
    "6) Observer — ReactiveX (StateFlow, INotifyPropertyChanged):\n"
    "   — UI подписывается на изменения состояния.\n\n"
    "7) Singleton — Django ORM (пул соединений с БД)."
)

add_qa_block(63,
    "Какие меры безопасности вы применили и какие планируете?",
    "Реализованные меры:\n"
    "• JWT-аутентификация с ограниченным временем жизни токенов;\n"
    "• Хеширование паролей (PBKDF2 + SHA-256);\n"
    "• CSRF-защита для session-based эндпоинтов (админка);\n"
    "• Разделение прав: IsAuthenticated для checkout и профиля.\n\n"
    "Планируемые (для production):\n"
    "• DEBUG = False — отключение режима отладки;\n"
    "• ALLOWED_HOSTS — конкретные домены вместо '*';\n"
    "• CORS_ALLOWED_ORIGINS — белый список вместо '*';\n"
    "• SECRET_KEY и пароли — в переменных окружения (.env);\n"
    "• Rate limiting — django-ratelimit или DRF throttling;\n"
    "• HTTPS — Let's Encrypt / Cloudflare;\n"
    "• Content Security Policy (CSP) — заголовки безопасности;\n"
    "• SQL injection — Django ORM (сырой SQL в seed.py — исключение);\n"
    "• XSS — React экранирует вывод (но HTML в description нужно "
    "санировать);\n"
    "• Logging — централизованный сбор логов;\n"
    "• Инвалидация refresh-токенов при logout."
)

add_qa_block(64,
    "Как вы тестировали проект?",
    "Тестирование выполнялось на разных уровнях:\n\n"
    "• Unit-тесты: 18 xUnit-тестов для десктопного плеера "
    "(PlayerViewModelTests.cs). Тестируют: создание плейлиста, "
    "поиск, режимы, дедупликацию, сериализацию.\n\n"
    "• Интеграционное тестирование Django:\n"
    "   — Проверка system check: python manage.py check (0 ошибок);\n"
    "   — Проверка API вручную через браузер (Browsable API DRF);\n"
    "   — Проверка seed-скрипта (запуск многократно — идемпотентность);\n"
    "   — Проверка миграции PostgreSQL → SQLite.\n\n"
    "• Ручное тестирование:\n"
    "   — Desktop плеер: воспроизведение, поиск, эквалайзер, "
    "сохранение/загрузка плейлиста;\n"
    "   — Android плеер: сканирование, воспроизведение, фоновый режим;\n"
    "   — Веб-фронтенд: все 9 страниц, корзина, аутентификация;\n"
    "   — React Native: навигация, drawer, мок-данные.\n\n"
    "• Сборка: Android APK (app-debug.apk, 18.3 MB).\n\n"
    "Не тестировалось: нагрузочное тестирование (stress test), "
    "тестирование на реальных устройствах (Android — только эмулятор), "
    "автоматизированное UI-тестирование (Selenium / Detox)."
)

add_qa_block(65,
    "Какие инструменты и IDE вы использовали?",
    "Инструменты разработки:\n"
    "• Visual Studio 2022 / Rider — разработка .NET (Desktop плеер);\n"
    "• Android Studio — разработка Android (Kotlin + Compose);\n"
    "• VS Code — разработка Python, TypeScript, React;\n"
    "• Python 3.11 + venv — виртуальное окружение Django;\n"
    "• Node.js 22 — Vite, React, React Native;\n"
    "• .NET SDK 10 — сборка десктопного плеера;\n"
    "• JDK 17.0.12 — сборка Android APK;\n"
    "• Android SDK 34 + Build Tools 34.0.0;\n"
    "• sdkmanager — установка Android SDK;\n"
    "• Git — система контроля версий;\n"
    "• PostgreSQL 16 / SQLite 3.45.1 — базы данных;\n"
    "• pgAdmin — управление PostgreSQL;\n"
    "• Adobe Photoshop / Figma — дизайн (презентация, иконки);\n"
    "• python-pptx / python-docx — генерация документов;\n"
    "• xUnit / .NET Test SDK — тестирование."
)

add_qa_block(66,
    "Какие версии зависимостей вы используете и почему?",
    "Ключевые версии и обоснование:\n\n"
    ".NET 10 — последняя версия .NET на момент разработки, LTS-поддержка, "
    "AOT-компиляция, улучшенная производительность GC.\n\n"
    "Avalonia UI 11.1 — стабильная версия с FluentTheme, AcrylicBlur, "
    "поддержкой .NET 10.\n\n"
    "Android Gradle Plugin 8.2.2 — совместимость с JDK 17, "
    "стабильная версия для compileSdk 34.\n\n"
    "Kotlin 1.9.22 — совместимость с Compose compiler 1.5.8 "
    "(для Compose BOM 2024.02.00).\n\n"
    "AndroidX Media3 1.2.1 — стабильная версия, заменила устаревшие "
    "MediaCompat и ExoPlayer в одном пакете.\n\n"
    "Django 5.2 — последняя LTS-версия (поддержка до 2027+).\n\n"
    "React 18.3.1 — стабильная версия с Concurrent Mode и "
    "автоматическим батчингом.\n\n"
    "Vite 5.2.11 — быстрый сборщик, совместим с React 18 и TypeScript 6.\n\n"
    "Expo SDK 56 — последняя стабильная версия на момент разработки, "
    "включает React Native 0.85.3."
)

add_qa_block(67,
    "Чем отличается структура проекта на разных платформах?",
    "Структура проекта отражает специфику каждой платформы:\n\n"
    "Десктоп (C# / Avalonia):\n"
    "• 2 проекта в Solution: Player (приложение) + Player.Tests (тесты);\n"
    "• MVVM-разделение: Models / ViewModels / Views;\n"
    "• Сервисы: EqualizerSampleProvider (DSP), MusicScanner (файловая система);\n"
    "• Точка входа: Program.cs → App.axaml → MainWindow.axaml.\n\n"
    "Android (Kotlin / Compose):\n"
    "• Пакеты: ui (ViewModel + Screen), service (PlaybackService), "
    "data (Track + MusicScanner), ui/theme (Color + Theme);\n"
    "• Single-Activity (MainActivity), Service (PlaybackService);\n"
    "• Точка входа: MainActivity.kt → setContent { PlayerTheme { ... } }.\n\n"
    "Бэкенд (Python / Django):\n"
    "• config/ — настройки (settings, urls, wsgi, asgi);\n"
    "• shop/ — приложение (models, views, serializers, admin, urls);\n"
    "• shop/admin2/ — кастомная админ-панель (12 HTML-шаблонов);\n"
    "• shop/management/commands/ — seed-скрипты.\n\n"
    "Веб-фронтенд (React / TypeScript):\n"
    "• src/pages/ — 9 страниц-компонентов;\n"
    "• src/components/ — переиспользуемые компоненты (Header, Marquee и др.);\n"
    "• src/context/ — AuthContext;\n"
    "• src/types/ — TypeScript-интерфейсы.\n\n"
    "React Native:\n"
    "• src/screens/ — 9 экранов, src/components/ — переиспользуемые, "
    "src/context/ — контексты, src/styles/ — тема, src/api/ — клиент."
)

add_qa_block(68,
    "Какие ошибки вы обнаружили при анализе кода?",
    "При анализе кодовой базы (analysis_report.txt) обнаружено:\n\n"
    "Критические:\n"
    "1) Android: PlaybackService не вызывает startForeground() → "
    "IllegalStateException на Android 13+ (краш при фоновом воспроизведении).\n"
    "2) Android: плейлист полностью пересоздаётся при каждом playTrack() "
    "(неэффективная работа с ExoPlayer).\n"
    "3) Django: DEBUG=True, ALLOWED_HOSTS=['*'], "
    "CORS_ALLOW_ALL_ORIGINS=True — опасно для продакшена.\n"
    "4) Django: SECRET_KEY и пароль БД в открытом виде.\n"
    "5) Django: UserProfile.__str__() — self.display_name или self.user.username, "
    "но user_id — IntegerField, не ForeignKey → упадёт, если user_id не существует.\n"
    "6) Django/Frontend: порт 8001 во фронтенде vs 8000 в бэкенде — несовместимость.\n\n"
    "Важные:\n"
    "7) Нет .gitignore для .env, архитектурных артефактов.\n"
    "8) Нет MEDIA_URL / MEDIA_ROOT для изображений.\n"
    "9) seed_pg.py использует сырой SQL (не Django ORM).\n"
    "10) ShopUser не наследуется от AbstractUser — две параллельные "
    "таблицы пользователей.\n"
    "11) Desktop: не чистый MVVM (ViewModel в code-behind).\n"
    "12) Обложки без кэширования.\n"
    "13) IsLoading / LoadingProgress никогда не устанавливаются.\n"
    "14) index.css — 1688 строк в одном файле.\n"
    "15) Нет адаптивной верстки."
)

add_qa_block(69,
    "Как вы обеспечиваете переносимость кода?",
    "Переносимость обеспечивается следующим:\n\n"
    "• Avalonia UI — кроссплатформенный UI-фреймворк, работает на "
    "Windows, macOS, Linux без изменения кода;\n"
    "• Android — стандартная сборка через Gradle, APK совместим "
    "с Android 8.0+ (API 26+);\n"
    "• Django — переносим между PostgreSQL, SQLite, MySQL (достаточно "
    "изменить ENGINE в settings.py);\n"
    "• React / Vite — стандартный сборщик, не привязан к платформе;\n"
    "• React Native / Expo — работает на Android и iOS (iOS требует Mac);\n"
    "• BAT-файлы — Windows-only; для Linux/Mac нужна адаптация (bash-скрипты);\n"
    "• Миграция БД PostgreSQL → SQLite — подтверждает возможность "
    "смены СУБД без потери данных.\n\n"
    "Код написан на стандартных API (Django ORM, React, Avalonia, "
    "AndroidX) без использования проприетарных/закрытых библиотек."
)

doc.add_page_break()

# ──────────────────────────────────────────────
# SECTION 9: Comparison & Future
# ──────────────────────────────────────────────
add_section_header("Сравнение платформ и перспективы развития", "Раздел 9")

add_qa_block(70,
    "Сравните десктопный и Android-плеер: плюсы и минусы.",
    "Десктопный плеер (C# / Avalonia):\n"
    "Плюсы: эквалайзер, ID3-теги + обложки, сохранение плейлиста, "
    "сканирование всех дисков, больше форматов, 18 unit-тестов.\n"
    "Минусы: нет фонового воспроизведения, привязан к окну, "
    "больше потребление памяти (~50-80 MB), Windows-only (теоретически "
    "кроссплатформенный, но тестировался только на Windows).\n\n"
    "Android-плеер (Kotlin / Compose):\n"
    "Плюсы: фоновое воспроизведение, интеграция с MediaStore, "
    "удобство мобильного формата, управление с экрана блокировки.\n"
    "Минусы: нет эквалайзера, нет обложек, нет сохранения плейлиста, "
    "обработка разрешений Android 13+, критический баг с startForeground().\n\n"
    "Общие черты: одинаковая цветовая палитра, общая логика "
    "(Repeat/Shuffle/AutoPlay, поиск, управление плейлистом), "
    "архитектура MVVM."
)

add_qa_block(71,
    "Сравните веб-фронтенд и React Native: какой подход эффективнее?",
    "Веб-фронтенд (React + Vite + TypeScript):\n"
    "Плюсы: быстрая разработка, HMR (горячая перезагрузка), "
    "богатая экосистема, отладка в браузере (DevTools), "
    "CSS-анимации, SEO (при SSR).\n"
    "Минусы: нет push-уведомлений, нет доступа к нативным API "
    "(камера, геолокация), офлайн работает только через Service Worker.\n\n"
    "React Native (Expo + TypeScript):\n"
    "Плюсы: нативный UI (FlatList вместо виртуализации браузера), "
    "доступ к нативным API (AsyncStorage, геолокация, камера), "
    "push-уведомления, работа в фоне.\n"
    "Минусы: дольше сборка, больше размер приложения (Expo ~20 MB), "
    "меньше готовых компонентов, сложнее отладка (React Native Debugger).\n\n"
    "Вывод: React Native не заменяет веб, а дополняет его. "
    "Веб — для десктопных пользователей, RN — для мобильных."
)

add_qa_block(72,
    "Какие функции вы бы добавили в первую очередь?",
    "Приоритетные улучшения:\n\n"
    "1) Исправление критического бага Android — добавить "
    "startForeground() в PlaybackService.\n\n"
    "2) Синхронизация портов — исправить порт 8001 на 8000 "
    "во фронтенде (App.tsx, AuthContext).\n\n"
    "3) Эквалайзер на Android — перенести DSP-фильтры "
    "из десктопной версии (BiQuadFilter).\n\n"
    "4) Сохранение плейлиста на Android — через Room или "
    "SharedPreferences (JSON).\n\n"
    "5) Кэширование обложек — на диск (MemoryCache + LRU).\n\n"
    "6) Адаптивная верстка — media queries для планшетов и телефонов.\n\n"
    "7) Модульные CSS — разбить index.css на CSS Modules.\n\n"
    "8) Docker — контейнеризация бэкенда и фронтенда.\n\n"
    "9) CI/CD — GitHub Actions для тестов и сборки.\n\n"
    "10) Деплой — разместить фронтенд на Vercel/Netlify, "
    "бэкенд на Railway/Render."
)

add_qa_block(73,
    "Как можно улучшить архитектуру проекта?",
    "Архитектурные улучшения:\n\n"
    "1) Разделение монолитного фронтенда — вынести общую логику "
    "(типы, API-клиент, утилиты) в отдельный пакет (shared/) "
    "для переиспользования между вебом и React Native.\n\n"
    "2) DI (Dependency Injection) в десктопном плеере — "
    "внедрить Microsoft.Extensions.DependencyInjection "
    "для сервисов (IAudioService, IPlaylistService, ITagReader).\n\n"
    "3) Правильные связи — IntegerField user_id заменить на ForeignKey "
    "(если перенести ShopUser в managed=True).\n\n"
    "4) Транзакции в checkout — обернуть создание заказа "
    "в transaction.atomic() для атомарности.\n\n"
    "5) F() для stock — использовать F('stock') - quantity "
    "для атомарного обновления стока.\n\n"
    "6) CustomUser — ShopUser наследует AbstractUser "
    "(вместо ручного дублирования полей).\n\n"
    "7) API Versioning — /api/v1/... для обратной совместимости.\n\n"
    "8) OpenAPI/Swagger — документирование API через drf-spectacular."
)

add_qa_block(74,
    "Как вы оцениваете практическую значимость проекта?",
    "Практическая значимость:\n\n"
    "1) Навыки — получен практический опыт работы с 5 языками "
    "(C#, Kotlin, Python, TypeScript, SQL), 4 платформами "
    "(Desktop, Android, Web, Mobile), 3 архитектурными паттернами "
    "(MVVM, MVC, Service Layer).\n\n"
    "2) Демонстрация full-stack разработки — проект охватывает "
    "весь стек: от low-level DSP-фильтров (NAudio BiQuadFilter) "
    "до высокоуровневого UI (React, Compose) и облачных сервисов "
    "(Django API, JWT).\n\n"
    "3) Решаемые проблемы:\n"
    "   — Плеер: нет качественных бесплатных аудиоплееров для "
    "Windows с тёмной темой и эквалайзером;\n"
    "   — Магазин: нишевый маркетплейс мерча российских "
    "рэп-исполнителей (аналог Lamoda / Ozon, но специализированный);\n\n"
    "4) Потенциальное применение:\n"
    "   — Кастомизация под конкретного исполнителя (white-label);\n"
    "   — Интеграция с мерч-службами (CDEK, Boxberry для расчёта доставки);\n"
    "   — Добавление онлайн-оплаты (ЮKassa, SberPay)."
)

add_qa_block(75,
    "Ваши выводы по результатам дипломного проекта?",
    "Выводы:\n\n"
    "1) Цели достигнуты — разработаны две версии аудиоплеера "
    "(десктоп и Android) и полнофункциональный интернет-магазин "
    "(веб + мобильное приложение + бэкенд + админка).\n\n"
    "2) Технологии подтвердили свою эффективность:\n"
    "   — Avalonia UI + NAudio — зрелые библиотеки для Desktop;\n"
    "   — Jetpack Compose + Media3 — современный стандарт Android;\n"
    "   — Django + DRF — быстрая разработка API;\n"
    "   — React + Vite — производительная SPA;\n"
    "   — React Native — быстрая мобильная разработка.\n\n"
    "3) Выявлены области для улучшения — 25+ проблем, "
    "от критических (startForeground, несовпадение портов) "
    "до косметических (один CSS-файл, мок-данные).\n\n"
    "4) Навыки, приобретённые в ходе работы:\n"
    "   — Разработка под 4 платформы;\n"
    "   — Работа с БД (миграция PostgreSQL ↔ SQLite);\n"
    "   — JWT-аутентификация;\n"
    "   — MVVM на C# и Kotlin;\n"
    "   — Аудиообработка (DSP-фильтры, MediaSession);\n"
    "   — Документирование (Python генерация .docx/.pptx).\n\n"
    "5) Проект готов к защите и может быть использован как "
    "основа для дальнейшего коммерческого развития."
)

# ══════════════════════════════════════════════════════════════
#                     SAVE
# ══════════════════════════════════════════════════════════════
output_path = "E:/илья проект/вопросы_и_ответы.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
